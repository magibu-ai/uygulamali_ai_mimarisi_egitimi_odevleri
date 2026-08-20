"""Safe, explicit loaders for common user-facing document formats."""

from __future__ import annotations

import csv
from pathlib import Path
from typing import Iterable

from dynamic_rag.models import Document

SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".xlsx", ".docx", ".pdf"}


def _load_text(path: Path) -> list[Document]:
    return [Document(path.read_text(encoding="utf-8", errors="replace"), path.name, path.stem)]


def _load_csv(path: Path) -> list[Document]:
    with path.open(encoding="utf-8-sig", errors="replace", newline="") as file:
        rows = list(csv.DictReader(file))
    return [
        Document(
            "\n".join(f"{key}: {value}" for key, value in row.items() if value),
            path.name,
            f"{path.stem} — satır {index}",
            {"row": index},
        )
        for index, row in enumerate(rows, start=2)
        if any(row.values())
    ]


def _load_xlsx(path: Path) -> list[Document]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    documents = []
    for sheet in workbook.worksheets:
        rows = sheet.iter_rows(values_only=True)
        headers = [str(value or f"column_{i}") for i, value in enumerate(next(rows, []), start=1)]
        for index, values in enumerate(rows, start=2):
            pairs = [f"{h}: {v}" for h, v in zip(headers, values) if v not in (None, "")]
            if pairs:
                documents.append(Document("\n".join(pairs), path.name, f"{sheet.title} — satır {index}", {"sheet": sheet.title, "row": index}))
    return documents


def _load_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    text = "\n\n".join(p.text.strip() for p in DocxDocument(path).paragraphs if p.text.strip())
    return [Document(text, path.name, path.stem)] if text else []


def _load_pdf(path: Path) -> list[Document]:
    from pypdf import PdfReader

    documents = []
    for page_number, page in enumerate(PdfReader(path).pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            documents.append(Document(text, path.name, f"{path.stem} — sayfa {page_number}", {"page": page_number}))
    return documents


LOADERS = {".txt": _load_text, ".md": _load_text, ".csv": _load_csv, ".xlsx": _load_xlsx, ".docx": _load_docx, ".pdf": _load_pdf}


def load_files(paths: Iterable[str | Path], raw_text: str = "") -> list[Document]:
    documents = []
    for item in paths:
        path = Path(item)
        suffix = path.suffix.lower()
        if suffix not in LOADERS:
            raise ValueError(f"Desteklenmeyen dosya türü: {suffix or '(uzantısız)'}")
        documents.extend(LOADERS[suffix](path))
    if raw_text.strip():
        documents.append(Document(raw_text.strip(), "manual-text", "Elle girilen metin"))
    if not documents:
        raise ValueError("En az bir okunabilir dosya veya metin gerekli.")
    return documents


def load_hf_dataset(repo_id: str, *, split: str = "train", text_column: str = "text", max_rows: int = 1000, token: str | None = None) -> list[Document]:
    from datasets import load_dataset

    dataset = load_dataset(repo_id, split=split, token=token)
    if text_column not in dataset.column_names:
        raise ValueError(f"Metin sütunu bulunamadı: {text_column}; mevcut={dataset.column_names}")
    documents = []
    for index, row in enumerate(dataset.select(range(min(max_rows, len(dataset))))):
        text = str(row.get(text_column) or "").strip()
        if text:
            documents.append(Document(text, f"hf://{repo_id}/{split}/{index}", str(row.get("title") or f"Kayıt {index}"), {"row": index, "repo_id": repo_id}))
    if not documents:
        raise ValueError("Seçilen Hugging Face veri setinde okunabilir metin yok.")
    return documents
